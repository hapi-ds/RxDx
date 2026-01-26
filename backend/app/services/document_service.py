"""
Document generation service for RxDx.

This service handles generation of various document types including:
- Design Review PDFs (ReportLab)
- Traceability Matrix PDFs (ReportLab)
- FMEA Excel files (openpyxl)
- Invoice Word documents (python-docx-template)

Implements Requirement 8 (Document Generation).
"""

import hashlib
import io
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.db.graph import GraphService
from app.models.user import User
from app.schemas.document import (
    DesignReviewRequest,
    DesignReviewResponse,
    DocumentFormat,
    DocumentRecord,
    DocumentStatus,
    DocumentType,
    FMEARequest,
    FMEAResponse,
    InvoiceLineItem,
    InvoiceRequest,
    InvoiceResponse,
    TraceabilityMatrixRequest,
    TraceabilityMatrixResponse,
)
from app.services.audit_service import AuditService
from app.services.signature_service import SignatureService

# Default template directory
TEMPLATE_DIR = Path(__file__).parent.parent / "templates"


class DocumentService:
    """
    Service for generating compliance documents.
    
    Supports PDF, Excel, and Word document generation with
    digital signature inclusion and audit logging.
    """

    def __init__(
        self,
        graph_service: GraphService,
        audit_service: AuditService,
        signature_service: SignatureService,
        template_dir: Path | None = None,
    ):
        """
        Initialize DocumentService.
        
        Args:
            graph_service: Service for graph database operations
            audit_service: Service for audit logging
            signature_service: Service for digital signature operations
            template_dir: Directory containing document templates
        """
        self.graph_service = graph_service
        self.audit_service = audit_service
        self.signature_service = signature_service
        self.template_dir = template_dir or TEMPLATE_DIR

        # Document storage (in production, use proper storage service)
        self._document_store: dict[UUID, DocumentRecord] = {}

    # ========================================================================
    # Helper Methods
    # ========================================================================

    def _get_styles(self) -> dict[str, ParagraphStyle]:
        """Get custom paragraph styles for PDF generation."""
        styles = getSampleStyleSheet()

        # Custom title style
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1a365d'),
        ))

        # Custom heading style
        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2c5282'),
        ))

        # Custom body style
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceBefore=6,
            spaceAfter=6,
        ))

        return styles

    def _create_header_table_style(self) -> TableStyle:
        """Create standard table style with header formatting."""
        return TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f7fafc')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ])

    def _calculate_content_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of document content."""
        return hashlib.sha256(content).hexdigest()

    async def _store_document(
        self,
        document_id: UUID,
        project_id: UUID,
        document_type: DocumentType,
        format: DocumentFormat,
        filename: str,
        content: bytes,
        user: User,
        metadata: dict[str, Any],
    ) -> DocumentRecord:
        """Store generated document and create record."""
        content_hash = self._calculate_content_hash(content)

        record = DocumentRecord(
            id=document_id,
            project_id=project_id,
            document_type=document_type,
            format=format,
            status=DocumentStatus.COMPLETED,
            filename=filename,
            file_path=f"/documents/{document_id}/{filename}",
            file_size_bytes=len(content),
            content_hash=content_hash,
            version="1.0",
            metadata=metadata,
            generated_at=datetime.now(UTC),
            generated_by=user.id,
        )

        # Store document (in production, use proper storage)
        self._document_store[document_id] = record

        return record

    async def _get_requirements_for_project(
        self,
        project_id: UUID,
        requirement_ids: list[UUID] | None = None,
    ) -> list[dict[str, Any]]:
        """Get requirements from graph database."""
        # Query graph for requirements
        requirements = await self.graph_service.get_workitems_by_type(
            str(project_id),
            'requirement'
        )

        # Filter by specific IDs if provided
        if requirement_ids:
            id_set = {str(rid) for rid in requirement_ids}
            requirements = [r for r in requirements if r.get('id') in id_set]

        return requirements

    async def _get_signatures_for_workitem(
        self,
        workitem_id: UUID,
    ) -> list[dict[str, Any]]:
        """Get signatures for a workitem."""
        signatures = await self.signature_service.get_workitem_signatures(workitem_id)
        return [
            {
                'signer_name': sig.user_name if hasattr(sig, 'user_name') else 'Unknown',
                'signed_at': sig.signed_at.isoformat() if sig.signed_at else '',
                'is_valid': sig.is_valid,
                'version': sig.workitem_version,
            }
            for sig in signatures
        ]

    # ========================================================================
    # Design Review PDF Generation
    # ========================================================================

    async def generate_design_review_pdf(
        self,
        request: DesignReviewRequest,
        user: User,
    ) -> DesignReviewResponse:
        """
        Generate a design phase review PDF document.
        
        Args:
            request: Design review generation request
            user: User generating the document
            
        Returns:
            Design review response with document details
        """
        document_id = uuid4()
        now = datetime.now(UTC)

        # Get requirements
        requirements = await self._get_requirements_for_project(
            request.project_id,
            request.requirement_ids,
        )

        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = self._get_styles()
        story = []

        # Title
        title = request.title or "Design Phase Review"
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Metadata section
        story.append(Paragraph(f"Project ID: {request.project_id}", styles['CustomBody']))
        story.append(Paragraph(f"Generated: {now.strftime('%Y-%m-%d %H:%M:%S UTC')}", styles['CustomBody']))
        story.append(Paragraph(f"Generated By: {user.full_name}", styles['CustomBody']))
        story.append(Paragraph(f"Total Requirements: {len(requirements)}", styles['CustomBody']))
        story.append(Spacer(1, 24))

        # Requirements section
        total_signatures = 0

        for i, req in enumerate(requirements, 1):
            # Requirement header
            req_title = req.get('title', 'Untitled Requirement')
            req_id = req.get('id', 'N/A')
            story.append(Paragraph(
                f"{i}. {req_title}",
                styles['CustomHeading']
            ))

            # Requirement details table
            details_data = [
                ['Field', 'Value'],
                ['ID', str(req_id)],
                ['Version', req.get('version', '1.0')],
                ['Status', req.get('status', 'draft')],
                ['Priority', str(req.get('priority', 'N/A'))],
            ]

            details_table = Table(details_data, colWidths=[2*inch, 4*inch])
            details_table.setStyle(self._create_header_table_style())
            story.append(details_table)
            story.append(Spacer(1, 12))

            # Description
            description = req.get('description', 'No description provided.')
            story.append(Paragraph("Description:", styles['CustomBody']))
            story.append(Paragraph(description, styles['CustomBody']))
            story.append(Spacer(1, 12))

            # Signatures section
            if request.include_signatures:
                try:
                    signatures = await self._get_signatures_for_workitem(UUID(str(req_id)))
                    if signatures:
                        total_signatures += len(signatures)
                        sig_data = [['Signer', 'Date', 'Version', 'Status']]
                        for sig in signatures:
                            sig_data.append([
                                sig['signer_name'],
                                sig['signed_at'][:19] if sig['signed_at'] else 'N/A',
                                sig['version'],
                                'Valid' if sig['is_valid'] else 'Invalid',
                            ])

                        sig_table = Table(sig_data, colWidths=[1.5*inch, 2*inch, 1*inch, 1*inch])
                        sig_table.setStyle(self._create_header_table_style())
                        story.append(Paragraph("Digital Signatures:", styles['CustomBody']))
                        story.append(sig_table)
                except Exception:
                    pass  # Skip signatures if error

            story.append(Spacer(1, 24))

        # Build PDF
        doc.build(story)
        content = buffer.getvalue()

        # Generate filename
        filename = f"design_review_{request.project_id}_{now.strftime('%Y%m%d_%H%M%S')}.pdf"

        # Store document
        await self._store_document(
            document_id=document_id,
            project_id=request.project_id,
            document_type=DocumentType.DESIGN_REVIEW,
            format=DocumentFormat.PDF,
            filename=filename,
            content=content,
            user=user,
            metadata={
                'requirement_count': len(requirements),
                'signature_count': total_signatures,
                'include_signatures': request.include_signatures,
            },
        )

        # Log audit event
        await self.audit_service.log(
            user_id=user.id,
            action="GENERATE",
            entity_type="Document",
            entity_id=document_id,
            details={
                'document_type': DocumentType.DESIGN_REVIEW.value,
                'project_id': str(request.project_id),
                'requirement_count': len(requirements),
            },
        )

        return DesignReviewResponse(
            document_id=document_id,
            project_id=request.project_id,
            status=DocumentStatus.COMPLETED,
            filename=filename,
            file_size_bytes=len(content),
            requirement_count=len(requirements),
            signature_count=total_signatures,
            generated_at=now,
            generated_by=user.id,
            download_url=f"/api/v1/documents/{document_id}",
        )

    # ========================================================================
    # Traceability Matrix PDF Generation
    # ========================================================================

    async def generate_traceability_matrix_pdf(
        self,
        request: TraceabilityMatrixRequest,
        user: User,
    ) -> TraceabilityMatrixResponse:
        """
        Generate a requirements traceability matrix PDF.
        
        Args:
            request: Traceability matrix generation request
            user: User generating the document
            
        Returns:
            Traceability matrix response with document details
        """
        document_id = uuid4()
        now = datetime.now(UTC)

        # Get traceability data from graph
        matrix_data = await self.graph_service.get_traceability_matrix(str(request.project_id))

        # Filter by specific requirement IDs if provided
        if request.requirement_ids:
            id_set = {str(rid) for rid in request.requirement_ids}
            matrix_data = [row for row in matrix_data if row.get('requirement_id') in id_set]

        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=36,
            leftMargin=36,
            topMargin=72,
            bottomMargin=72,
        )

        styles = self._get_styles()
        story = []

        # Title
        story.append(Paragraph("Requirements Traceability Matrix", styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Metadata
        story.append(Paragraph(f"Project ID: {request.project_id}", styles['CustomBody']))
        story.append(Paragraph(f"Generated: {now.strftime('%Y-%m-%d %H:%M:%S UTC')}", styles['CustomBody']))
        story.append(Paragraph(f"Generated By: {user.full_name}", styles['CustomBody']))
        story.append(Spacer(1, 24))

        # Build table data
        headers = ['Requirement']
        if request.include_tests:
            headers.append('Tests')
        if request.include_risks:
            headers.append('Risks')
        if request.include_signatures:
            headers.append('Signed')

        table_data = [headers]

        test_count = 0
        risk_count = 0
        covered_count = 0

        for row in matrix_data:
            row_data = [row.get('requirement_title', 'Unknown')[:50]]

            if request.include_tests:
                tests = row.get('test_ids', [])
                test_count += len(tests)
                if tests:
                    covered_count += 1
                row_data.append(', '.join(tests[:3]) + ('...' if len(tests) > 3 else '') if tests else 'None')

            if request.include_risks:
                risks = row.get('risk_ids', [])
                risk_count += len(risks)
                row_data.append(', '.join(risks[:3]) + ('...' if len(risks) > 3 else '') if risks else 'None')

            if request.include_signatures:
                is_signed = row.get('is_signed', False)
                row_data.append('Yes' if is_signed else 'No')

            table_data.append(row_data)

        # Calculate column widths
        num_cols = len(headers)
        col_width = (A4[0] - 72) / num_cols
        col_widths = [col_width] * num_cols

        # Create table
        table = Table(table_data, colWidths=col_widths)
        table.setStyle(self._create_header_table_style())
        story.append(table)
        story.append(Spacer(1, 24))

        # Summary section
        total_requirements = len(matrix_data)
        coverage = (covered_count / total_requirements * 100) if total_requirements > 0 else 0

        story.append(Paragraph("Summary", styles['CustomHeading']))
        summary_data = [
            ['Metric', 'Value'],
            ['Total Requirements', str(total_requirements)],
            ['Requirements with Tests', str(covered_count)],
            ['Test Coverage', f"{coverage:.1f}%"],
            ['Total Tests', str(test_count)],
            ['Total Risks', str(risk_count)],
        ]
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(self._create_header_table_style())
        story.append(summary_table)

        # Build PDF
        doc.build(story)
        content = buffer.getvalue()

        # Generate filename
        filename = f"traceability_matrix_{request.project_id}_{now.strftime('%Y%m%d_%H%M%S')}.pdf"

        # Store document
        await self._store_document(
            document_id=document_id,
            project_id=request.project_id,
            document_type=DocumentType.TRACEABILITY_MATRIX,
            format=DocumentFormat.PDF,
            filename=filename,
            content=content,
            user=user,
            metadata={
                'requirement_count': total_requirements,
                'test_count': test_count,
                'risk_count': risk_count,
                'coverage_percentage': coverage,
            },
        )

        # Log audit event
        await self.audit_service.log(
            user_id=user.id,
            action="GENERATE",
            entity_type="Document",
            entity_id=document_id,
            details={
                'document_type': DocumentType.TRACEABILITY_MATRIX.value,
                'project_id': str(request.project_id),
                'requirement_count': total_requirements,
            },
        )

        return TraceabilityMatrixResponse(
            document_id=document_id,
            project_id=request.project_id,
            status=DocumentStatus.COMPLETED,
            filename=filename,
            file_size_bytes=len(content),
            requirement_count=total_requirements,
            test_count=test_count,
            risk_count=risk_count,
            coverage_percentage=coverage,
            generated_at=now,
            generated_by=user.id,
            download_url=f"/api/v1/documents/{document_id}",
        )

    # ========================================================================
    # FMEA Excel Generation
    # ========================================================================

    async def generate_fmea_excel(
        self,
        request: FMEARequest,
        user: User,
    ) -> FMEAResponse:
        """
        Generate an FMEA Excel document with risk chains.
        
        Args:
            request: FMEA generation request
            user: User generating the document
            
        Returns:
            FMEA response with document details
        """
        document_id = uuid4()
        now = datetime.now(UTC)

        # Get all risks from graph
        risks = await self.graph_service.get_all_risks(str(request.project_id))

        # Filter by specific risk IDs if provided
        if request.risk_ids:
            id_set = {str(rid) for rid in request.risk_ids}
            risks = [r for r in risks if r.get('id') in id_set]

        # Filter by minimum RPN if provided
        if request.min_rpn:
            risks = [r for r in risks if r.get('rpn', 0) >= request.min_rpn]

        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "FMEA Analysis"

        # Define styles
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2C5282", end_color="2C5282", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        cell_alignment = Alignment(vertical="top", wrap_text=True)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin'),
        )

        # Headers
        headers = [
            'Risk ID',
            'Title',
            'Description',
            'Failure Mode',
            'Failure Effect',
            'Failure Cause',
            'Severity (S)',
            'Occurrence (O)',
            'Detection (D)',
            'RPN',
            'Risk Level',
            'Current Controls',
        ]

        if request.include_failure_chains:
            headers.append('Failure Chain')

        if request.include_mitigations:
            headers.extend(['Mitigation Actions', 'Mitigation Status'])

        # Write headers
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # Set column widths
        column_widths = {
            'A': 15,  # Risk ID
            'B': 25,  # Title
            'C': 40,  # Description
            'D': 25,  # Failure Mode
            'E': 25,  # Failure Effect
            'F': 25,  # Failure Cause
            'G': 12,  # Severity
            'H': 12,  # Occurrence
            'I': 12,  # Detection
            'J': 10,  # RPN
            'K': 12,  # Risk Level
            'L': 30,  # Current Controls
            'M': 40,  # Failure Chain
            'N': 40,  # Mitigation Actions
            'O': 15,  # Mitigation Status
        }

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # Track statistics
        failure_count = 0
        mitigation_count = 0
        total_rpn = 0
        max_rpn = 0

        # Write risk data
        for row_idx, risk in enumerate(risks, 2):
            rpn = risk.get('rpn', 0)
            total_rpn += rpn
            max_rpn = max(max_rpn, rpn)

            # Determine risk level
            if rpn >= 200:
                risk_level = "Critical"
            elif rpn >= 100:
                risk_level = "High"
            elif rpn >= 50:
                risk_level = "Medium"
            else:
                risk_level = "Low"

            row_data = [
                str(risk.get('id', ''))[:8],
                risk.get('title', ''),
                risk.get('description', ''),
                risk.get('failure_mode', ''),
                risk.get('failure_effect', ''),
                risk.get('failure_cause', ''),
                risk.get('severity', 0),
                risk.get('occurrence', 0),
                risk.get('detection', 0),
                rpn,
                risk_level,
                risk.get('current_controls', ''),
            ]

            # Get failure chains if requested
            if request.include_failure_chains:
                try:
                    chains = await self.graph_service.get_risk_chains(risk.get('id'))
                    if chains:
                        failure_count += len(chains)
                        chain_str = ' -> '.join([
                            f.get('description', 'Unknown')[:30]
                            for f in chains[0].get('path', [])[:5]
                        ])
                        row_data.append(chain_str)
                    else:
                        row_data.append('')
                except Exception:
                    row_data.append('')

            # Get mitigations if requested
            if request.include_mitigations:
                mitigations = risk.get('mitigations', [])
                mitigation_count += len(mitigations)

                if mitigations:
                    mitigation_titles = ', '.join([
                        m.get('title', 'Unknown')[:30]
                        for m in mitigations[:3]
                    ])
                    mitigation_statuses = ', '.join([
                        m.get('status', 'unknown')
                        for m in mitigations[:3]
                    ])
                    row_data.extend([mitigation_titles, mitigation_statuses])
                else:
                    row_data.extend(['None', 'N/A'])

            # Write row
            for col, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.alignment = cell_alignment
                cell.border = thin_border

                # Color code RPN column
                if col == 10:  # RPN column
                    if value >= 200:
                        cell.fill = PatternFill(start_color="FC8181", end_color="FC8181", fill_type="solid")
                    elif value >= 100:
                        cell.fill = PatternFill(start_color="F6AD55", end_color="F6AD55", fill_type="solid")
                    elif value >= 50:
                        cell.fill = PatternFill(start_color="FAF089", end_color="FAF089", fill_type="solid")
                    else:
                        cell.fill = PatternFill(start_color="9AE6B4", end_color="9AE6B4", fill_type="solid")

        # Freeze header row
        ws.freeze_panes = 'A2'

        # Add summary sheet
        summary_ws = wb.create_sheet(title="Summary")
        summary_data = [
            ['FMEA Summary Report', ''],
            ['', ''],
            ['Project ID', str(request.project_id)],
            ['Generated', now.strftime('%Y-%m-%d %H:%M:%S UTC')],
            ['Generated By', user.full_name],
            ['', ''],
            ['Statistics', ''],
            ['Total Risks', len(risks)],
            ['Total Failures', failure_count],
            ['Total Mitigations', mitigation_count],
            ['Average RPN', f"{total_rpn / len(risks):.1f}" if risks else "0"],
            ['Maximum RPN', max_rpn],
        ]

        for row_idx, (label, value) in enumerate(summary_data, 1):
            summary_ws.cell(row=row_idx, column=1, value=label)
            summary_ws.cell(row=row_idx, column=2, value=value)

        summary_ws.column_dimensions['A'].width = 20
        summary_ws.column_dimensions['B'].width = 40

        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        content = buffer.getvalue()

        # Generate filename
        filename = f"fmea_{request.project_id}_{now.strftime('%Y%m%d_%H%M%S')}.xlsx"

        # Store document
        await self._store_document(
            document_id=document_id,
            project_id=request.project_id,
            document_type=DocumentType.FMEA,
            format=DocumentFormat.EXCEL,
            filename=filename,
            content=content,
            user=user,
            metadata={
                'risk_count': len(risks),
                'failure_count': failure_count,
                'mitigation_count': mitigation_count,
                'average_rpn': total_rpn / len(risks) if risks else 0,
                'max_rpn': max_rpn,
            },
        )

        # Log audit event
        await self.audit_service.log(
            user_id=user.id,
            action="GENERATE",
            entity_type="Document",
            entity_id=document_id,
            details={
                'document_type': DocumentType.FMEA.value,
                'project_id': str(request.project_id),
                'risk_count': len(risks),
            },
        )

        return FMEAResponse(
            document_id=document_id,
            project_id=request.project_id,
            status=DocumentStatus.COMPLETED,
            filename=filename,
            file_size_bytes=len(content),
            risk_count=len(risks),
            failure_count=failure_count,
            mitigation_count=mitigation_count,
            average_rpn=total_rpn / len(risks) if risks else 0,
            max_rpn=max_rpn,
            generated_at=now,
            generated_by=user.id,
            download_url=f"/api/v1/documents/{document_id}",
        )

    # ========================================================================
    # Invoice Word Document Generation
    # ========================================================================

    async def generate_invoice_word(
        self,
        request: InvoiceRequest,
        user: User,
    ) -> InvoiceResponse:
        """
        Generate an invoice Word document using a template.
        
        Args:
            request: Invoice generation request
            user: User generating the document
            
        Returns:
            Invoice response with document details
        """
        document_id = uuid4()
        now = datetime.now(UTC)

        # Get time entries for the billing period
        time_entries = await self._get_time_entries_for_period(
            project_id=request.project_id,
            start_date=request.billing_period.start_date,
            end_date=request.billing_period.end_date,
        )

        # Aggregate time entries by user and task
        aggregated: dict[tuple, dict[str, Any]] = {}
        for entry in time_entries:
            key = (entry.get('user_id'), entry.get('task_id'))
            if key not in aggregated:
                aggregated[key] = {
                    'user_name': entry.get('user_name', 'Unknown'),
                    'task_title': entry.get('task_title', 'Unknown Task'),
                    'hours': 0.0,
                    'rate': entry.get('hourly_rate', 0.0),
                }
            aggregated[key]['hours'] += entry.get('duration_hours', 0.0)

        # Calculate totals
        line_items: list[InvoiceLineItem] = []
        subtotal = 0.0
        total_hours = 0.0

        for item in aggregated.values():
            item_subtotal = item['hours'] * item['rate']
            line_items.append(InvoiceLineItem(
                description=f"{item['user_name']} - {item['task_title']}",
                hours=item['hours'],
                rate=item['rate'],
                subtotal=item_subtotal,
            ))
            subtotal += item_subtotal
            total_hours += item['hours']

        # Calculate tax and total
        tax_amount = subtotal * request.tax_rate
        total_amount = subtotal + tax_amount

        # Try to load template, or create simple document
        try:
            template_path = self.template_dir / f"{request.template_name}.docx"
            if template_path.exists():
                doc = DocxTemplate(str(template_path))
            else:
                # Create a simple document without template
                content = self._create_simple_invoice_content(
                    request=request,
                    line_items=line_items,
                    subtotal=subtotal,
                    tax_amount=tax_amount,
                    total_amount=total_amount,
                    total_hours=total_hours,
                    user=user,
                    now=now,
                )
                filename = f"invoice_{request.project_id}_{now.strftime('%Y%m%d_%H%M%S')}.docx"

                await self._store_document(
                    document_id=document_id,
                    project_id=request.project_id,
                    document_type=DocumentType.INVOICE,
                    format=DocumentFormat.WORD,
                    filename=filename,
                    content=content,
                    user=user,
                    metadata={
                        'billing_period_start': request.billing_period.start_date.isoformat(),
                        'billing_period_end': request.billing_period.end_date.isoformat(),
                        'total_hours': total_hours,
                        'subtotal': subtotal,
                        'tax_amount': tax_amount,
                        'total_amount': total_amount,
                    },
                )

                await self.audit_service.log(
                    user_id=user.id,
                    action="GENERATE",
                    entity_type="Document",
                    entity_id=document_id,
                    details={
                        'document_type': DocumentType.INVOICE.value,
                        'project_id': str(request.project_id),
                        'total_amount': total_amount,
                    },
                )

                return InvoiceResponse(
                    document_id=document_id,
                    project_id=request.project_id,
                    status=DocumentStatus.COMPLETED,
                    filename=filename,
                    file_size_bytes=len(content),
                    billing_period_start=request.billing_period.start_date,
                    billing_period_end=request.billing_period.end_date,
                    total_hours=total_hours,
                    subtotal=subtotal,
                    tax_amount=tax_amount,
                    total_amount=total_amount,
                    line_item_count=len(line_items),
                    generated_at=now,
                    generated_by=user.id,
                    download_url=f"/api/v1/documents/{document_id}",
                )
        except Exception:
            # Fall back to simple document
            content = self._create_simple_invoice_content(
                request=request,
                line_items=line_items,
                subtotal=subtotal,
                tax_amount=tax_amount,
                total_amount=total_amount,
                total_hours=total_hours,
                user=user,
                now=now,
            )
            filename = f"invoice_{request.project_id}_{now.strftime('%Y%m%d_%H%M%S')}.docx"

            await self._store_document(
                document_id=document_id,
                project_id=request.project_id,
                document_type=DocumentType.INVOICE,
                format=DocumentFormat.WORD,
                filename=filename,
                content=content,
                user=user,
                metadata={
                    'billing_period_start': request.billing_period.start_date.isoformat(),
                    'billing_period_end': request.billing_period.end_date.isoformat(),
                    'total_hours': total_hours,
                    'subtotal': subtotal,
                    'tax_amount': tax_amount,
                    'total_amount': total_amount,
                },
            )

            await self.audit_service.log(
                user_id=user.id,
                action="GENERATE",
                entity_type="Document",
                entity_id=document_id,
                details={
                    'document_type': DocumentType.INVOICE.value,
                    'project_id': str(request.project_id),
                    'total_amount': total_amount,
                },
            )

            return InvoiceResponse(
                document_id=document_id,
                project_id=request.project_id,
                status=DocumentStatus.COMPLETED,
                filename=filename,
                file_size_bytes=len(content),
                billing_period_start=request.billing_period.start_date,
                billing_period_end=request.billing_period.end_date,
                total_hours=total_hours,
                subtotal=subtotal,
                tax_amount=tax_amount,
                total_amount=total_amount,
                line_item_count=len(line_items),
                generated_at=now,
                generated_by=user.id,
                download_url=f"/api/v1/documents/{document_id}",
            )

        # Render template with context
        context = {
            'invoice_number': f"INV-{now.strftime('%Y%m%d')}-{str(document_id)[:8].upper()}",
            'invoice_date': now.strftime('%Y-%m-%d'),
            'billing_period': f"{request.billing_period.start_date} to {request.billing_period.end_date}",
            'client_name': request.client_name or 'Client',
            'client_address': request.client_address or '',
            'project_id': str(request.project_id),
            'line_items': [item.model_dump() for item in line_items],
            'subtotal': f"{subtotal:.2f}",
            'tax_rate': f"{request.tax_rate * 100:.0f}%",
            'tax_amount': f"{tax_amount:.2f}",
            'total': f"{total_amount:.2f}",
            'currency': request.currency,
            'notes': request.notes or '',
            'generated_by': user.full_name,
        }
        doc.render(context)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        # Generate filename
        filename = f"invoice_{request.project_id}_{now.strftime('%Y%m%d_%H%M%S')}.docx"

        # Store document
        await self._store_document(
            document_id=document_id,
            project_id=request.project_id,
            document_type=DocumentType.INVOICE,
            format=DocumentFormat.WORD,
            filename=filename,
            content=content,
            user=user,
            metadata={
                'billing_period_start': request.billing_period.start_date.isoformat(),
                'billing_period_end': request.billing_period.end_date.isoformat(),
                'total_hours': total_hours,
                'subtotal': subtotal,
                'tax_amount': tax_amount,
                'total_amount': total_amount,
            },
        )

        # Log audit event
        await self.audit_service.log(
            user_id=user.id,
            action="GENERATE",
            entity_type="Document",
            entity_id=document_id,
            details={
                'document_type': DocumentType.INVOICE.value,
                'project_id': str(request.project_id),
                'total_amount': total_amount,
            },
        )

        return InvoiceResponse(
            document_id=document_id,
            project_id=request.project_id,
            status=DocumentStatus.COMPLETED,
            filename=filename,
            file_size_bytes=len(content),
            billing_period_start=request.billing_period.start_date,
            billing_period_end=request.billing_period.end_date,
            total_hours=total_hours,
            subtotal=subtotal,
            tax_amount=tax_amount,
            total_amount=total_amount,
            line_item_count=len(line_items),
            generated_at=now,
            generated_by=user.id,
            download_url=f"/api/v1/documents/{document_id}",
        )

    def _create_simple_invoice_content(
        self,
        request: InvoiceRequest,
        line_items: list[InvoiceLineItem],
        subtotal: float,
        tax_amount: float,
        total_amount: float,
        total_hours: float,
        user: User,
        now: datetime,
    ) -> bytes:
        """Create a simple invoice document without template."""
        doc = DocxDocument()

        # Title
        title = doc.add_heading('INVOICE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Invoice details
        doc.add_paragraph(f"Invoice Number: INV-{now.strftime('%Y%m%d')}-{str(uuid4())[:8].upper()}")
        doc.add_paragraph(f"Date: {now.strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Billing Period: {request.billing_period.start_date} to {request.billing_period.end_date}")
        doc.add_paragraph()

        # Client info
        if request.client_name:
            doc.add_paragraph(f"Bill To: {request.client_name}")
        if request.client_address:
            doc.add_paragraph(request.client_address)
        doc.add_paragraph()

        # Line items table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Description'
        header_cells[1].text = 'Hours'
        header_cells[2].text = 'Rate'
        header_cells[3].text = 'Amount'

        # Data rows
        for item in line_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.description
            row_cells[1].text = f"{item.hours:.2f}"
            row_cells[2].text = f"{request.currency} {item.rate:.2f}"
            row_cells[3].text = f"{request.currency} {item.subtotal:.2f}"

        doc.add_paragraph()

        # Totals
        doc.add_paragraph(f"Subtotal: {request.currency} {subtotal:.2f}")
        doc.add_paragraph(f"Tax ({request.tax_rate * 100:.0f}%): {request.currency} {tax_amount:.2f}")
        doc.add_paragraph(f"Total: {request.currency} {total_amount:.2f}")

        # Notes
        if request.notes:
            doc.add_paragraph()
            doc.add_paragraph(f"Notes: {request.notes}")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    async def _get_time_entries_for_period(
        self,
        project_id: UUID,
        start_date,
        end_date,
    ) -> list[dict[str, Any]]:
        """Get time entries for a billing period."""
        # In production, this would query the time entries from the database
        # For now, return empty list as time service is not yet implemented
        return []

    # ========================================================================
    # Document Retrieval
    # ========================================================================

    async def get_document(
        self,
        document_id: UUID,
        user: User,
    ) -> DocumentRecord | None:
        """
        Retrieve a document record by ID.
        
        Args:
            document_id: Document ID
            user: User requesting the document
            
        Returns:
            Document record if found, None otherwise
        """
        record = self._document_store.get(document_id)

        if record:
            # Log access
            await self.audit_service.log(
                user_id=user.id,
                action="READ",
                entity_type="Document",
                entity_id=document_id,
                details={'filename': record.filename},
            )

        return record

    async def get_document_content(
        self,
        document_id: UUID,
        user: User,
    ) -> bytes | None:
        """
        Retrieve document content by ID.
        
        Args:
            document_id: Document ID
            user: User requesting the document
            
        Returns:
            Document content bytes if found, None otherwise
        """
        # In production, this would retrieve from storage service
        # For now, return None as we don't persist content in memory
        record = await self.get_document(document_id, user)
        if not record:
            return None

        # In production implementation, retrieve from file storage
        return None

    async def list_documents(
        self,
        project_id: UUID | None = None,
        document_type: DocumentType | None = None,
        limit: int = 100,
        offset: int = 0,
    ) -> list[DocumentRecord]:
        """
        List documents with optional filtering.
        
        Args:
            project_id: Filter by project ID
            document_type: Filter by document type
            limit: Maximum number of results
            offset: Number of results to skip
            
        Returns:
            List of document records
        """
        records = list(self._document_store.values())

        # Apply filters
        if project_id:
            records = [r for r in records if r.project_id == project_id]

        if document_type:
            records = [r for r in records if r.document_type == document_type]

        # Sort by generated_at descending
        records.sort(key=lambda r: r.generated_at, reverse=True)

        # Apply pagination
        return records[offset:offset + limit]


# ============================================================================
# Dependency Injection
# ============================================================================

def get_document_service(
    graph_service: GraphService,
    audit_service: AuditService,
    signature_service: SignatureService,
) -> DocumentService:
    """
    Factory function for DocumentService dependency injection.
    
    Args:
        graph_service: Graph database service
        audit_service: Audit logging service
        signature_service: Digital signature service
        
    Returns:
        DocumentService instance
    """
    return DocumentService(
        graph_service=graph_service,
        audit_service=audit_service,
        signature_service=signature_service,
    )
